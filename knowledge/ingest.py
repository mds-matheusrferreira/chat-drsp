#!/usr/bin/env python3
import argparse
import csv
import json
import os
import sys
from pathlib import Path


def ensure_chromadb_python():
    # PATCH PARA XAMPP/APACHE: Garante as variáveis de ambiente necessárias
    if os.name == 'nt':
        default_user = r"C:\Users\avisala.cebas"
        username = "avisala.cebas"
        
        if 'USERPROFILE' not in os.environ:
            os.environ['USERPROFILE'] = default_user
        if 'HOME' not in os.environ:
            os.environ['HOME'] = default_user
        if 'USERNAME' not in os.environ:
            os.environ['USERNAME'] = username
        if 'USER' not in os.environ:
            os.environ['USER'] = username
    try:
        import chromadb
        from chromadb.utils import embedding_functions

        return chromadb, embedding_functions
    except ModuleNotFoundError:
        fallback = os.environ.get('KNOWLEDGE_PYTHON_FALLBACK', 'C:/Python314/python.exe')
        current = Path(sys.executable).resolve()
        fallback_path = Path(fallback)

        if os.name == 'nt':
            user_site = 'C:/Users/avisala.cebas/AppData/Roaming/Python/Python314/site-packages'
            os.environ['PYTHONPATH'] = user_site + os.pathsep + os.environ.get('PYTHONPATH', '')

        if os.name == 'nt' and fallback_path.exists() and current != fallback_path.resolve():
            os.execv(str(fallback_path), [str(fallback_path), *sys.argv])

        raise


chromadb, embedding_functions = ensure_chromadb_python()

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


SUPPORTED_EXTENSIONS = {'.txt', '.md', '.csv', '.pdf', '.docx', '.xlsx', '.xls'}


def project_root() -> Path:
    return Path(__file__).resolve().parents[1]


def chroma_path() -> Path:
    configured = os.environ.get('KNOWLEDGE_CHROMA_PATH')
    if configured:
        path = Path(configured)
        return path if path.is_absolute() else project_root() / path

    return project_root() / 'storage/app/private/knowledge/chromadb'


def collection():
    if chromadb is None or embedding_functions is None:
        raise RuntimeError('Dependência chromadb não instalada.')

    path = chroma_path()
    path.mkdir(parents=True, exist_ok=True)

    client = chromadb.PersistentClient(path=str(path))
    embedder = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=os.environ.get('KNOWLEDGE_EMBEDDING_MODEL', 'all-MiniLM-L6-v2')
    )

    return client.get_or_create_collection(
        name=os.environ.get('KNOWLEDGE_COLLECTION', 'drsp_knowledge'),
        embedding_function=embedder,
        metadata={'description': 'Base interna DRSP/SUAS'},
    )


def extract_text(path: Path) -> str:
    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f'Tipo de arquivo não suportado: {extension}')

    if extension in {'.txt', '.md'}:
        return path.read_text(encoding='utf-8', errors='ignore')

    if extension == '.csv':
        rows = []
        with path.open('r', encoding='utf-8', errors='ignore', newline='') as file:
            reader = csv.reader(file)
            for row in reader:
                rows.append(' | '.join(cell.strip() for cell in row if cell is not None))
        return '\n'.join(rows)

    if extension == '.pdf':
        if PdfReader is None:
            raise RuntimeError('Dependência pypdf não instalada.')

        reader = PdfReader(str(path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ''
            if text.strip():
                pages.append(f'[Página {index}]\n{text}')
        return '\n\n'.join(pages)

    if extension == '.docx':
        if DocxDocument is None:
            raise RuntimeError('Dependência python-docx não instalada.')

        document = DocxDocument(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(' | '.join(values))

        return '\n'.join(parts)

    if extension in {'.xlsx', '.xls'}:
        if extension == '.xlsx' and load_workbook is not None:
            # Uso de Context Manager (with) para liberar a memória do arquivo após leitura
            with load_workbook(filename=str(path), read_only=True, data_only=True) as workbook:
                lines = []
                for sheet in workbook.worksheets:
                    lines.append(f'[Aba: {sheet.title}]')
                    for row in sheet.iter_rows(values_only=True):
                        values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                        if values:
                            lines.append(' | '.join(values))
                return '\n'.join(lines)

        if pd is None:
            raise RuntimeError('Dependência pandas/openpyxl não instalada.')

        sheets = pd.read_excel(str(path), sheet_name=None)
        lines = []
        for sheet_name, frame in sheets.items():
            lines.append(f'[Aba: {sheet_name}]')
            lines.append(frame.fillna('').to_csv(index=False, sep='|'))
        return '\n'.join(lines)

    return ''


def env_int(name: str, default: int) -> int:
    value = os.environ.get(name)

    if value is None or value == '':
        return default

    return int(value)


def validate_chunking(chunk_size: int, overlap: int) -> None:
    if chunk_size < 200:
        raise ValueError('chunk-size deve ser pelo menos 200.')

    if overlap < 0:
        raise ValueError('chunk-overlap não pode ser negativo.')

    if overlap >= chunk_size:
        raise ValueError('chunk-overlap deve ser menor que chunk-size.')


def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120) -> list[str]:
    # Normalização de quebras de linha redundantes
    normalized = '\n'.join(line.strip() for line in text.splitlines())
    normalized = '\n'.join(line for line in normalized.splitlines() if line)

    if not normalized:
        return []

    chunks = []
    start = 0
    text_len = len(normalized)

    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = normalized[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end == text_len:
            break

        # Evita loops infinitos e calcula o deslocamento correto com base no overlap
        step = chunk_size - overlap
        if step <= 0:
            step = 1  # Fallback de segurança mecânica
            
        start += step

    return chunks


def main():
    parser = argparse.ArgumentParser(description='Indexa documento interno DRSP no ChromaDB.')
    parser.add_argument('--document-id', required=True)
    parser.add_argument('--path', required=True)
    parser.add_argument('--title', required=True)
    parser.add_argument('--original-name', default='')
    parser.add_argument('--chunk-size', type=int, default=env_int('KNOWLEDGE_CHUNK_SIZE', 700))
    parser.add_argument('--chunk-overlap', type=int, default=env_int('KNOWLEDGE_CHUNK_OVERLAP', 120))
    args = parser.parse_args()

    validate_chunking(args.chunk_size, args.chunk_overlap)

    file_path = Path(args.path)

    if not file_path.exists():
        raise FileNotFoundError(f'Arquivo não encontrado: {file_path}')

    text = extract_text(file_path)
    chunks = chunk_text(text, chunk_size=args.chunk_size, overlap=args.chunk_overlap)

    if not chunks:
        raise ValueError('Não foi possível extrair texto útil do documento.')

    ids = [f'doc-{args.document_id}-chunk-{index}' for index in range(len(chunks))]
    metadatas = [
        {
            'document_id': str(args.document_id),
            'title': args.title,
            'original_name': args.original_name,
            'chunk_index': index,
            'extension': file_path.suffix.lower().lstrip('.'),
        }
        for index in range(len(chunks))
    ]

    store = collection()
    # Limpa vetores antigos desse ID antes de inserir os novos blocos modificados
    store.delete(where={'document_id': str(args.document_id)})
    store.add(ids=ids, documents=chunks, metadatas=metadatas)

    print(json.dumps({
        'status': 'ready',
        'chunks_count': len(chunks),
        'characters_count': len(text),
        'chunk_size': args.chunk_size,
        'chunk_overlap': args.chunk_overlap,
    }, ensure_ascii=False))


if __name__ == '__main__':
    try:
        main()
    except Exception as error:
        print(json.dumps({
            'status': 'failed',
            'error': str(error),
        }, ensure_ascii=False))
        raise SystemExit(1)